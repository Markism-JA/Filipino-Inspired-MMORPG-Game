import argparse
import os
import subprocess
import tempfile
from datetime import datetime
from docx import Document


def robust_replace_in_paragraph(paragraph, placeholder, replacement):
    count = 0
    runs = paragraph.runs
    if not runs:
        return 0

    full_text = "".join(run.text for run in runs)
    if placeholder not in full_text:
        return 0

    full_text = full_text.replace(placeholder, replacement)
    count = 1

    for i, run in enumerate(runs):
        run.text = ""
    runs[0].text = full_text

    return count


def replace_in_document(doc, placeholder, replacement):
    count = 0
    for p in doc.paragraphs:
        count += robust_replace_in_paragraph(p, placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    count += robust_replace_in_paragraph(p, placeholder, replacement)

    for section in doc.sections:
        for p in section.header.paragraphs:
            count += robust_replace_in_paragraph(p, placeholder, replacement)
        for p in section.footer.paragraphs:
            count += robust_replace_in_paragraph(p, placeholder, replacement)

    return count


def convert_to_pdf(input_docx, output_dir):
    print(f"Converting {input_docx} to PDF in {output_dir}...")
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                input_docx,
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except FileNotFoundError:
        print("Error: 'libreoffice' command not found.")
        print("Please install LibreOffice to enable PDF conversion.")
        return False
    except subprocess.CalledProcessError as e:
        print(f"Error during PDF conversion: {e.stderr}")
        return False
    except subprocess.TimeoutExpired:
        print("Error: PDF conversion timed out after 60 seconds.")
        return False
    return True


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_file = os.path.join(
        script_dir, "Accomplishment-and-Consultation-Form.docx"
    )

    parser = argparse.ArgumentParser(
        description="Replace placeholders in a .docx template and save as PDF."
    )
    parser.add_argument("output", help="Output .pdf file path")

    parser.add_argument(
        "--week", required=True, help="The week number to insert (e.g., '5')"
    )
    parser.add_argument(
        "--placeholder",
        default="{{DATE}}",
        help="Date placeholder text to replace (default: {{DATE}})",
    )
    parser.add_argument(
        "--week-placeholder",
        default="{{WEEK}}",
        help="Week number placeholder text to replace (default: {{WEEK}})",
    )
    parser.add_argument(
        "--date", help="Date string to insert (default: today's date in mm/dd/YYYY)"
    )
    args = parser.parse_args()

    try:
        doc = Document(template_file)
    except Exception as e:
        print(f"Error opening template {template_file}: {e}")
        return

    date_replacement = args.date or datetime.now().strftime("%m/%d/%Y")
    week_replacement = args.week

    count_date = replace_in_document(doc, args.placeholder, date_replacement)
    count_week = replace_in_document(doc, args.week_placeholder, week_replacement)

    print(
        f"Replacing {args.placeholder!r} -> {date_replacement!r} ({count_date} found)"
    )
    print(
        f"Replacing {args.week_placeholder!r} -> {week_replacement!r} ({count_week} found)"
    )

    temp_docx = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_docx_path = f.name
        doc.save(temp_docx_path)

        final_pdf_path = os.path.abspath(args.output)
        final_pdf_dir = os.path.dirname(final_pdf_path)

        os.makedirs(final_pdf_dir, exist_ok=True)

        if convert_to_pdf(temp_docx_path, final_pdf_dir):
            converted_pdf_name = (
                os.path.splitext(os.path.basename(temp_docx_path))[0] + ".pdf"
            )
            converted_pdf_path = os.path.join(final_pdf_dir, converted_pdf_name)

            os.rename(converted_pdf_path, final_pdf_path)

            print(f"\nSuccessfully generated PDF: {final_pdf_path}")

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    finally:
        if temp_docx_path and os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)


if __name__ == "__main__":
    main()
